"""
docx_template.py — reusable helpers for generating formatted Chinese
technical documents.

使用方式:

    from docx_template import (
        new_doc, add_title_page, add_toc, add_heading, add_para,
        add_bullet, add_code, add_table, save,
    )

    doc = new_doc()
    add_title_page(
        doc,
        title='示例文档',
        subtitle='这是一个副标题',
        meta_lines=['文档日期：2026 年 4 月 15 日', '作者：xxx'],
    )
    add_toc(doc)

    add_heading(doc, '一、引言', level=1)
    add_para(doc, '这是一段正文……')
    add_bullet(doc, '要点 1')
    add_bullet(doc, '要点 2')

    add_heading(doc, '1.1 子章节', level=2)
    add_code(doc, 'aws sts get-caller-identity')
    add_table(doc, [
        ['列 A', '列 B', '列 C'],
        ['a1',   'b1',   'c1'],
        ['a2',   'b2',   'c2'],
    ])

    save(doc, '/path/to/output.docx')

格式规范详见同目录下 FORMAT_GUIDE.md。
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================================================
# 常量(改这里可以调整整个文档的视觉风格)
# ============================================================

FONT_LATIN = 'Calibri'
FONT_CJK = '微软雅黑'
FONT_CODE = 'Consolas'

COLOR_HEADING = RGBColor(0x1A, 0x36, 0x5D)       # 深蓝,一级到四级标题、表头背景
COLOR_SUBTITLE = RGBColor(0x55, 0x55, 0x77)      # 副标题
COLOR_META = RGBColor(0x66, 0x66, 0x66)          # 日期等元信息
COLOR_CODE = RGBColor(0x33, 0x33, 0x33)          # 代码块文字
COLOR_HEADER_BG = '1A365D'                        # 表头背景(hex,不带#)
COLOR_HEADER_TEXT = RGBColor(0xFF, 0xFF, 0xFF)   # 表头文字

HEADING_SIZES = {1: 18, 2: 15, 3: 12.5, 4: 11.5}

BODY_LINE_SPACING = 1.5
BODY_SPACE_AFTER = Pt(8)
HEADING_SPACE_BEFORE = Pt(18)
HEADING_SPACE_AFTER = Pt(10)
HEADING_LINE_SPACING = 1.3
BULLET_LINE_SPACING = 1.4
BULLET_SPACE_AFTER = Pt(4)
CODE_LINE_SPACING = 1.25


# ============================================================
# 内部工具
# ============================================================

def _set_cjk_font(run):
    """确保 run 的中文字符使用统一的东亚字体。"""
    if run.element.rPr is None:
        run.element.get_or_add_rPr()
    rFonts = run.element.rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        run.element.rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_CJK)


def _apply_base_styles(doc):
    """应用全局基础样式到 Normal / Heading / List Bullet 等内置样式。"""
    s = doc.styles['Normal']
    s.font.name = FONT_LATIN
    s.font.size = Pt(11)
    s.paragraph_format.space_after = BODY_SPACE_AFTER
    s.paragraph_format.line_spacing = BODY_LINE_SPACING
    rFonts = s.element.rPr.rFonts
    rFonts.set(qn('w:eastAsia'), FONT_CJK)

    for level, size in HEADING_SIZES.items():
        try:
            h = doc.styles[f'Heading {level}']
            h.font.name = FONT_LATIN
            h.font.size = Pt(size)
            h.font.bold = True
            h.font.color.rgb = COLOR_HEADING
            h.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CJK)
            h.paragraph_format.space_before = HEADING_SPACE_BEFORE
            h.paragraph_format.space_after = HEADING_SPACE_AFTER
            h.paragraph_format.line_spacing = HEADING_LINE_SPACING
        except KeyError:
            pass

    try:
        lb = doc.styles['List Bullet']
        lb.font.name = FONT_LATIN
        lb.font.size = Pt(11)
        lb.paragraph_format.space_after = BULLET_SPACE_AFTER
        lb.paragraph_format.line_spacing = BULLET_LINE_SPACING
        lb.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CJK)
    except KeyError:
        pass


# ============================================================
# 对外接口
# ============================================================

def new_doc():
    """创建一个空白文档并应用全局样式。"""
    doc = Document()
    _apply_base_styles(doc)
    return doc


def add_title_page(doc, title, subtitle=None, meta_lines=None):
    """添加标题页:大标题 + 副标题 + 元信息 + 分页。

    Args:
        doc: Document 对象
        title: 大标题文字
        subtitle: 副标题文字(可选)
        meta_lines: 元信息列表,每行一条(如日期、作者、版本)
    """
    doc.add_paragraph()
    doc.add_paragraph()

    # 大标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    p.paragraph_format.space_after = Pt(20)
    run = p.add_run(title)
    run.font.name = FONT_LATIN
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = COLOR_HEADING
    _set_cjk_font(run)

    # 副标题
    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(40)
        run = p.add_run(subtitle)
        run.font.name = FONT_LATIN
        run.font.size = Pt(18)
        run.font.color.rgb = COLOR_SUBTITLE
        _set_cjk_font(run)

    # 元信息
    for line in meta_lines or []:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(line)
        run.font.name = FONT_LATIN
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_META
        _set_cjk_font(run)

    doc.add_page_break()


def add_toc(doc, toc_heading='目录'):
    """添加一个可跳转的目录页。

    目录基于 Heading 1-3 自动生成。首次打开文档时 Word 会提示更新域,
    或用户可按 F9 / 右键 → 更新域 手动触发。
    """
    add_heading(doc, toc_heading, level=1)

    p = doc.add_paragraph()
    run = p.add_run()
    r = run._r

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    fldChar1.set(qn('w:dirty'), 'true')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = r'TOC \o "1-3" \h \z \u'

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    placeholder = OxmlElement('w:t')
    placeholder.text = '在 Word 中按 F9 或右键 → 更新域,以生成目录。'

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')

    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(placeholder)
    r.append(fldChar3)

    # 让 Word 打开文档时自动提示更新域
    settings = doc.settings.element
    updateFields = settings.find(qn('w:updateFields'))
    if updateFields is None:
        updateFields = OxmlElement('w:updateFields')
        updateFields.set(qn('w:val'), 'true')
        settings.append(updateFields)

    doc.add_page_break()


def add_heading(doc, text, level=1, page_break_before=None):
    """添加章节标题。

    Args:
        doc: Document 对象
        text: 标题文字
        level: 1-4
        page_break_before: 是否在标题前分页。
            None = 默认行为:一级标题自动分页、二级及以下不分页
            True/False = 强制指定
    """
    level = max(1, min(level, 4))

    # 决定是否分页
    if page_break_before is None:
        page_break_before = (level == 1)
    if page_break_before:
        last = doc.paragraphs[-1] if doc.paragraphs else None
        if last and last.text.strip():
            doc.add_page_break()

    p = doc.add_paragraph(style=f'Heading {level}')
    run = p.add_run(text)
    run.font.name = FONT_LATIN
    run.font.bold = True
    run.font.size = Pt(HEADING_SIZES[level])
    run.font.color.rgb = COLOR_HEADING
    _set_cjk_font(run)
    return p


def add_para(doc, text, bold=False, italic=False):
    """添加一段正文。"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT_LATIN
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic
    _set_cjk_font(run)
    return p


def add_bullet(doc, text, level=0, bold_prefix=None):
    """添加一个 bullet 列表项。

    Args:
        text: bullet 文字
        level: 缩进层级(0 / 1 / 2)
        bold_prefix: 如果提供,在 text 前加一段加粗前缀(如 "方向 1:")
    """
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after = BULLET_SPACE_AFTER
    p.paragraph_format.line_spacing = BULLET_LINE_SPACING

    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.font.name = FONT_LATIN
        run.font.size = Pt(11)
        run.bold = True
        _set_cjk_font(run)

    run = p.add_run(text)
    run.font.name = FONT_LATIN
    run.font.size = Pt(11)
    _set_cjk_font(run)
    return p


def add_code(doc, text):
    """添加一段代码块(等宽字体、浅缩进)。"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.line_spacing = CODE_LINE_SPACING
    run = p.add_run(text)
    run.font.name = FONT_CODE
    run.font.size = Pt(9.5)
    run.font.color.rgb = COLOR_CODE
    return p


def add_table(doc, rows, col_widths=None):
    """添加一个表格(第一行为表头,深蓝底白字加粗)。

    Args:
        rows: 2D 列表,rows[0] 是表头
        col_widths: 可选,每列宽度(英寸列表)
    """
    if not rows:
        return None
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Inches(w)

    # 表头
    for ci, header_text in enumerate(rows[0]):
        cell = table.rows[0].cells[ci]
        cell.text = header_text
        for p in cell.paragraphs:
            p.paragraph_format.line_spacing = HEADING_LINE_SPACING
            for r in p.runs:
                r.bold = True
                r.font.name = FONT_LATIN
                r.font.size = Pt(10)
                r.font.color.rgb = COLOR_HEADER_TEXT
                _set_cjk_font(r)
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), COLOR_HEADER_BG)
        tcPr.append(shd)

    # 数据行
    for ri, row in enumerate(rows[1:], start=1):
        for ci, val in enumerate(row):
            cell = table.rows[ri].cells[ci]
            cell.text = val
            for p in cell.paragraphs:
                p.paragraph_format.line_spacing = HEADING_LINE_SPACING
                p.paragraph_format.space_after = Pt(2)
                for r in p.runs:
                    r.font.name = FONT_LATIN
                    r.font.size = Pt(10)
                    _set_cjk_font(r)

    # 表格后留一点空白,避免下一段紧挨表格
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    return table


def save(doc, path):
    """保存文档。"""
    doc.save(path)
    print(f'Saved: {path}')


# ============================================================
# 使用示例(直接运行本文件即可生成一个示例文档)
# ============================================================

if __name__ == '__main__':
    doc = new_doc()

    add_title_page(
        doc,
        title='示例文档',
        subtitle='演示统一格式的所有元素',
        meta_lines=[
            '文档日期:2026 年 4 月 15 日',
            '适用范围:Bota 技术文档模板',
        ],
    )

    add_toc(doc)

    add_heading(doc, '一、引言', level=1)
    add_para(doc, '本示例文档演示了模板支持的所有元素:标题、副标题、正文、'
                  '加粗、斜体、列表、代码、表格。')

    add_heading(doc, '1.1 段落与列表', level=2)
    add_para(doc, '这是普通正文。')
    add_para(doc, '这是加粗正文。', bold=True)
    add_bullet(doc, '第一条要点')
    add_bullet(doc, '第二条要点')
    add_bullet(doc, '带粗体前缀的要点,用于强调类别或类型。', bold_prefix='要点 3:')

    add_heading(doc, '二、代码与表格', level=1)

    add_heading(doc, '2.1 代码块', level=2)
    add_para(doc, '示例命令:')
    add_code(doc, 'aws --profile dev logs filter-log-events \\\n'
                  '    --log-group-name /eks/bota-api \\\n'
                  '    --filter-pattern "ERROR"')

    add_heading(doc, '2.2 表格', level=2)
    add_table(doc, [
        ['字段', '类型', '说明'],
        ['id', 'string', '记录的唯一标识'],
        ['createdAt', 'timestamp', '创建时间(UTC)'],
        ['status', 'enum', 'pending / processing / done'],
    ])

    add_heading(doc, '三、结语', level=1)
    add_para(doc, '这份模板涵盖了常见的技术文档元素。新文档按这个模板生成'
                  '即可保证视觉统一。')

    save(doc, '_example_output.docx')
